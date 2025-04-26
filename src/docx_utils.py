import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_formatted_paragraph(doc, text):
    """Adds a paragraph to the document with proper formatting."""
    try:
        if text.startswith("### "):
            return doc.add_heading(text[4:].strip(), level=3)
        elif text.startswith("#### "):
            return doc.add_heading(text[5:].strip(), level=4)
        else:
            para = doc.add_paragraph()
            run = para.add_run()
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                else:
                    para.add_run(part)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            return para
    except Exception as e:
        logging.error(f"Failed to add formatted paragraph: {str(e)}", exc_info=True)
        raise

def set_margins(doc, margin):
    """Set document margins to the specified size in inches."""
    try:
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(margin)
            section.bottom_margin = Inches(margin)
            section.left_margin = Inches(margin)
            section.right_margin = Inches(margin)
    except Exception as e:
        logging.error(f"Failed to set document margins: {str(e)}", exc_info=True)
        raise

def save_to_docx(text, filepath, title):
    try:
        logging.info(f"Saving document to {filepath}")
        doc = Document()
        doc.add_heading(title, level=0)
        
        # Set margins (1.27 cm is approximately 0.5 inches)
        set_margins(doc, 0.5)  # 0.5 inches is approximately 1.27 cm

        lines = text.split("\n")
        for line in lines:
            if line.strip():  # Skip empty lines
                add_formatted_paragraph(doc, line.strip())
        
        doc.save(filepath)
        logging.info(f"Document saved successfully to {filepath}")
    except Exception as e:
        logging.error(f"Failed to save to docx: {str(e)}", exc_info=True)
        raise RuntimeError(f"Failed to save to docx: {str(e)}")
